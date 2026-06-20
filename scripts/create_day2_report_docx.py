from __future__ import annotations

import csv
import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
OUTPUT = REPORTS / "NYC_Taxi_Deliverable_2_Model_Optimization_Report.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)
BLACK = RGBColor(0, 0, 0)
HEADER_FILL = "F2F4F7"
BORDER = "A6A6A6"


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as file:
        return list(csv.DictReader(file))


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_table_cell_margins(table, top=80, bottom=80, start=120, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin_name, value in {
        "top": top,
        "bottom": bottom,
        "start": start,
        "end": end,
    }.items():
        element = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if element is None:
            element = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    width = tbl_pr.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), "9360")
    width.set(qn("w:type"), "dxa")

    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)

    width_dxa = [int(round(width_in * 1440)) for width_in in widths]
    for dxa in width_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(dxa))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    set_table_cell_margins(table)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(str(text))
    set_run_font(run, size=9.5, bold=bold)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    header_run = header.add_run("NYC Taxi Long-Trip Model | Deliverable 2")
    set_run_font(header_run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    footer_run = footer.add_run("Model Optimization and Ethical Evaluation")
    set_run_font(footer_run, size=9, color=MUTED)


def add_para(doc: Document, text: str, first_line=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    if first_line:
        paragraph.paragraph_format.first_line_indent = Inches(0.3)
    run = paragraph.add_run(text)
    set_run_font(run, size=11)
    return paragraph


def add_heading(doc: Document, text: str, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_title_page(doc: Document):
    for _ in range(2):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Deliverable 2: Model Optimization, Real-World Impact, and Ethical Evaluation")
    set_run_font(run, size=20, color=BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Predicting Long NYC Yellow Taxi Trips")
    set_run_font(run, size=13, color=MUTED, italic=True)

    metadata = [
        ("Project", "Real-World AI-Driven Analytics Solution Using Python"),
        ("Repository Branch", "project_day_2"),
        ("Prepared for", "Model optimization and ethical evaluation deliverable"),
        ("Prepared by", "Soumith Kumar Ananthula, Varun Reddy Beesam, Vivek Doppalapudi"),
        ("Date", date.today().strftime("%B %d, %Y")),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1.7, 4.8])
    for row_idx, (label, value) in enumerate(metadata):
        label_cell, value_cell = table.rows[row_idx].cells
        set_cell_shading(label_cell, HEADER_FILL)
        set_cell_text(label_cell, label, bold=True)
        set_cell_text(value_cell, value)

    doc.add_page_break()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_text(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.RIGHT if idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, align=align)

    set_table_geometry(table, widths)
    doc.add_paragraph()


def pct_delta(new: float, old: float) -> str:
    return f"{(new - old) * 100:+.2f} pts"


def format_class_weight(value: object) -> str:
    if not isinstance(value, dict):
        return str(value)
    try:
        items = sorted(value.items(), key=lambda item: int(item[0]))
    except (TypeError, ValueError):
        items = sorted(value.items(), key=lambda item: str(item[0]))
    return "{" + ", ".join(f"{key}: {weight}" for key, weight in items) + "}"


def build_document():
    metrics = read_json(REPORTS / "metrics.json")
    model_rows = read_csv(REPORTS / "model_metrics.csv")
    tuning_rows = read_csv(REPORTS / "tuning_results.csv")

    baseline = metrics["baseline_most_frequent"]
    initial = metrics["initial_logistic_regression"]
    final = metrics["optimized_hist_gradient_boosting"]
    selected = metrics["model_optimization"]["selected_candidate"]
    selected_tuning = next(row for row in tuning_rows if row["selected"] == "True")
    class_weight_text = format_class_weight(selected["class_weight"])

    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    add_heading(doc, "Objective", level=1)
    add_para(
        doc,
        "The objective of this deliverable is to refine the NYC taxi long-trip prediction model, tune hyperparameters, evaluate the final outcomes, and reflect on ethical, practical, and future considerations. The model predicts whether a cleaned yellow taxi trip is likely to last at least 30 minutes. This second version improves the Day 1 workflow by adding a validation set, comparing optimized model candidates, tuning the decision threshold, and documenting how those choices changed model behavior.",
    )

    add_heading(doc, "1. Hyperparameter Tuning", level=1)
    add_para(
        doc,
        "The Day 1 model used a class-balanced logistic regression classifier. That model was useful as an interpretable baseline, but it produced many false long-trip alerts. Its held-out precision for the 30-minute-plus class was 0.5183, which means only a little over half of predicted long trips were actually long. For Day 2, the model was refined to improve precision and overall accuracy while keeping recall high enough for operational planning.",
    )
    add_para(
        doc,
        f"The updated pipeline uses a 60/20/20 stratified train-validation-test split. The training set contains {metrics['train_test_split']['train_rows']:,} rows, the validation set contains {metrics['train_test_split']['validation_rows']:,} rows, and the final test set contains {metrics['train_test_split']['test_rows']:,} rows. The validation set was important because hyperparameters and the probability threshold were selected before the final test set was evaluated.",
    )
    add_para(
        doc,
        f"The optimized model family is histogram gradient boosting. Five candidate configurations were evaluated by changing max_iter, learning_rate, max_leaf_nodes, l2_regularization, class_weight, and the final classification threshold. The selected candidate was {selected['candidate']}, with max_iter = {selected['max_iter']}, learning_rate = {selected['learning_rate']}, max_leaf_nodes = {selected['max_leaf_nodes']}, l2_regularization = {selected['l2_regularization']}, and class_weight = {class_weight_text}. The final decision threshold was {metrics['model_optimization']['decision_threshold']:.2f}. Threshold selection used {metrics['model_optimization']['selection_metric']}, so precision was prioritized while requiring recall to stay at least as strong as the Day 1 logistic reference.",
    )

    add_table(
        doc,
        ["Selected tuning result", "Value"],
        [
            ["Validation accuracy", f"{float(selected_tuning['validation_accuracy']):.4f}"],
            ["Validation precision", f"{float(selected_tuning['validation_precision']):.4f}"],
            ["Validation recall", f"{float(selected_tuning['validation_recall']):.4f}"],
            ["Validation F0.5", f"{float(selected_tuning['validation_f0_5']):.4f}"],
            ["Selected threshold", f"{float(selected_tuning['selected_threshold']):.2f}"],
        ],
        [3.25, 3.25],
    )

    add_heading(doc, "2. Final Model Evaluation", level=1)
    add_para(
        doc,
        f"The final test set contained {metrics['train_test_split']['test_rows']:,} trips and preserved the long-trip rate of {metrics['train_test_split']['test_long_trip_rate'] * 100:.2f}%. Because long trips are the minority class, accuracy alone is not enough. A most-frequent baseline achieved {baseline['accuracy']:.4f} accuracy by predicting every trip as under 30 minutes, but it had 0.0000 precision, recall, and F1 for the long-trip class.",
    )
    add_para(
        doc,
        f"The optimized model improved the Day 1 model in the areas requested for Deliverable 2. Accuracy increased from {initial['accuracy']:.4f} to {final['accuracy']:.4f}, precision increased from {initial['precision']:.4f} to {final['precision']:.4f}, recall increased from {initial['recall']:.4f} to {final['recall']:.4f}, F1 increased from {initial['f1']:.4f} to {final['f1']:.4f}, and ROC-AUC increased from {initial['roc_auc']:.4f} to {final['roc_auc']:.4f}. The largest practical gain was precision: false positives dropped from {initial['false_positive']:,} to {final['false_positive']:,}. This means the optimized system produces far fewer unnecessary long-trip warnings.",
    )
    add_para(
        doc,
        f"The earlier strict-threshold version improved precision but did not protect long-trip recall, so the final tuning rule was revised. The logistic model caught {initial['recall'] * 100:.2f}% of actual long trips, while the optimized model now catches {final['recall'] * 100:.2f}%. The selected threshold of {final['decision_threshold']:.2f} keeps long-trip recall slightly above the logistic reference while still reducing unnecessary long-trip warnings.",
    )

    metric_rows = []
    for row in model_rows:
        metric_rows.append(
            [
                row["model"],
                f"{float(row['accuracy']):.4f}",
                f"{float(row['precision']):.4f}",
                f"{float(row['recall']):.4f}",
                f"{float(row['f1']):.4f}",
                f"{float(row['roc_auc']):.4f}",
            ]
        )
    add_table(
        doc,
        ["Model", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC"],
        metric_rows,
        [2.15, 0.85, 0.9, 0.85, 0.75, 1.0],
    )

    add_table(
        doc,
        ["Metric", "Day 1 logistic", "Day 2 optimized", "Change"],
        [
            ["Accuracy", f"{initial['accuracy']:.4f}", f"{final['accuracy']:.4f}", pct_delta(final["accuracy"], initial["accuracy"])],
            ["Precision", f"{initial['precision']:.4f}", f"{final['precision']:.4f}", pct_delta(final["precision"], initial["precision"])],
            ["Recall", f"{initial['recall']:.4f}", f"{final['recall']:.4f}", pct_delta(final["recall"], initial["recall"])],
            ["F1", f"{initial['f1']:.4f}", f"{final['f1']:.4f}", pct_delta(final["f1"], initial["f1"])],
            ["ROC-AUC", f"{initial['roc_auc']:.4f}", f"{final['roc_auc']:.4f}", pct_delta(final["roc_auc"], initial["roc_auc"])],
        ],
        [1.4, 1.55, 1.65, 1.9],
    )

    add_para(
        doc,
        f"The final confusion matrix also shows the effect of the changes. The optimized model produced {final['true_positive']:,} true positives, {final['true_negative']:,} true negatives, {final['false_positive']:,} false positives, and {final['false_negative']:,} false negatives. Compared with the logistic model, the optimized model reduced false positives by {initial['false_positive'] - final['false_positive']:,} and reduced false negatives by {initial['false_negative'] - final['false_negative']:,}. This means the final model is both less noisy and slightly better at catching actual long trips.",
    )

    add_heading(doc, "3. Ethical Considerations", level=1)
    add_para(
        doc,
        "The model does not use direct personal identifiers such as rider names, phone numbers, exact addresses, race, gender, or payment identity. However, ethical risk still exists because pickup and dropoff location IDs can reflect neighborhood patterns, airport access, commuting behavior, income differences, and historical service availability. A model trained on historical taxi records can reproduce existing transportation patterns rather than represent fair service goals.",
    )
    add_para(
        doc,
        "The Day 2 precision improvement reduces unnecessary long-trip alerts, but it does not remove fairness concerns. Higher precision means the model is less likely to incorrectly flag a short trip as long, which is helpful for planning. The revised threshold also protects recall, so long-trip coverage stays slightly above the logistic reference. Even so, if remaining errors are concentrated in certain zones or times, the model could still create uneven operational attention. Before using this system in production, the team should audit false positives and false negatives by pickup zone, dropoff zone, hour, weekday, weekend, and airport-trip status.",
    )
    add_para(
        doc,
        "The model should not be used to deny rides, discourage drivers from serving specific neighborhoods, punish drivers, or set prices without additional fairness review. The safest ethical use is aggregate planning with human oversight. Users should see the prediction as a probability-based estimate rather than a guaranteed statement about an individual trip.",
    )

    add_heading(doc, "4. Real-World Application", level=1)
    add_para(
        doc,
        "A practical deployment would be a dashboard for taxi dispatchers, fleet managers, airport queue monitors, or city transportation analysts. The dashboard could estimate the share of upcoming trips likely to last at least 30 minutes and show where long-trip pressure may affect vehicle availability. This is useful because long trips remove vehicles from high-demand zones for longer periods, which can influence staffing, airport staging, driver breaks, and customer wait-time estimates.",
    )
    add_para(
        doc,
        f"The Day 2 changes make the model more useful for this real-world setting because the optimized model is less noisy without sacrificing long-trip coverage. The logistic model produced many false positives, which could cause dispatchers to overreact. The optimized model raises precision to {final['precision']:.4f}, so a long-trip alert is more trustworthy. Accuracy also increased to {final['accuracy']:.4f}, recall increased to {final['recall']:.4f}, and ROC-AUC increased to {final['roc_auc']:.4f}, meaning the model separates long and short trips more effectively across probability thresholds.",
    )
    add_para(
        doc,
        f"The system should still show probability bands rather than only a hard class label. For example, trips could be grouped as low, medium, or high long-trip risk. The tuned {final['decision_threshold']:.2f} threshold can be used for the main operational long-trip flag because it balances better precision with recall that stays above the logistic reference. Other thresholds can still support monitoring dashboards when stakeholders want to favor either extra caution or fewer alerts. This gives stakeholders flexibility without retraining the model every time business priorities change.",
    )
    add_para(
        doc,
        "Future real-world improvements should include more months of TLC data, weather, public events, traffic incidents, construction schedules, and holiday indicators. Probability calibration should also be tested before probabilities are shown directly to decision-makers. Finally, the team should monitor model drift because travel behavior changes across seasons, major events, policy changes, and economic conditions.",
    )

    add_heading(doc, "5. Final Thoughts and Conclusion", level=1)
    add_para(
        doc,
        "Deliverable 2 successfully refined the Day 1 model into a stronger optimized classifier. The main technical changes were the addition of a validation set, the use of histogram gradient boosting, systematic hyperparameter tuning, validation-based threshold selection, and final model feature-importance reporting. These changes made the model more accurate and more precise while keeping long-trip recall protected.",
    )
    add_para(
        doc,
        f"The final model achieved {final['accuracy'] * 100:.2f}% accuracy, {final['precision'] * 100:.2f}% precision, {final['recall'] * 100:.2f}% recall, {final['f1'] * 100:.2f}% F1, and {final['roc_auc'] * 100:.2f}% ROC-AUC on the held-out test set. After retuning, the result is better than the Day 1 logistic reference across the main metrics, including recall. The important tradeoff is now threshold flexibility: a stricter threshold can raise precision further, but the submitted threshold is the better final choice because it improves precision while keeping recall above the logistic reference.",
    )
    add_para(
        doc,
        "The model is best used as a planning aid, not an automated decision system. It can help transportation teams understand long-trip patterns and prepare resources, but ethical use requires transparency, subgroup error audits, and human judgment. With those safeguards, the optimized model provides a stronger foundation for real-world taxi operations analytics.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
