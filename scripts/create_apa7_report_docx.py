from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
OUTPUT = REPORTS / "NYC_Taxi_AI_Analytics_APA7_Report.docx"


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as file:
        return list(csv.DictReader(file))


def set_cell_text(cell, text: str, bold: bool = False, align=None, font_size=10):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(text))
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_table_borders(table, color="000000", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.color.rgb = None
        style.paragraph_format.line_spacing = 2
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

    styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Heading 2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_page_number(section.header.paragraphs[0])


def add_paragraph(doc: Document, text: str, first_line=True):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 2
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return paragraph


def add_heading(doc: Document, text: str, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.line_spacing = 2
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = True
    return paragraph


def add_table_title(doc: Document, number: int, title: str):
    p_num = doc.add_paragraph()
    p_num.paragraph_format.line_spacing = 2
    p_num.paragraph_format.space_after = Pt(0)
    r_num = p_num.add_run(f"Table {number}")
    r_num.bold = True
    r_num.font.name = "Times New Roman"
    r_num.font.size = Pt(12)

    p_title = doc.add_paragraph()
    p_title.paragraph_format.line_spacing = 2
    p_title.paragraph_format.space_after = Pt(0)
    r_title = p_title.add_run(title)
    r_title.italic = True
    r_title.font.name = "Times New Roman"
    r_title.font.size = Pt(12)


def add_apa_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.style = "Table Grid"
    set_table_borders(table, color="666666", size="4")
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        set_cell_text(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=10)
        set_cell_shading(cell, "F2F2F2")

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 and value.replace(".", "", 1).replace(",", "").isdigit() else WD_ALIGN_PARAGRAPH.LEFT
            cells[i].width = Inches(widths[i])
            set_cell_text(cells[i], value, align=align, font_size=9)

    doc.add_paragraph()


def add_figure(doc: Document, number: int, title: str, image_path: Path, width=5.8):
    p_num = doc.add_paragraph()
    p_num.paragraph_format.line_spacing = 2
    r_num = p_num.add_run(f"Figure {number}")
    r_num.bold = True
    r_num.font.name = "Times New Roman"
    r_num.font.size = Pt(12)

    p_title = doc.add_paragraph()
    p_title.paragraph_format.line_spacing = 2
    r_title = p_title.add_run(title)
    r_title.italic = True
    r_title.font.name = "Times New Roman"
    r_title.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))


def add_references(doc: Document, references: list[str]):
    add_heading(doc, "References", level=1)
    for ref in references:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.line_spacing = 2
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
        run = paragraph.add_run(ref)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def build_document():
    metrics = read_json(REPORTS / "metrics.json")
    audit = read_csv(REPORTS / "preprocessing_audit.csv")
    features = read_csv(REPORTS / "feature_engineering_summary.csv")

    doc = Document()
    configure_document(doc)

    # APA 7 student title page.
    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 2
    title_run = title.add_run("Predicting Long NYC Yellow Taxi Trips Using Python-Based AI Analytics")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(12)

    for line in [
        "Soumith Kumar Ananthula",
        "Varun Reddy Beesam",
        "Vivek Doppalapudi",
        "Real-World AI-Driven Analytics Solution Using Python",
        "Project Day 1",
        "June 19, 2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2
        r = p.add_run(line)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

    doc.add_page_break()
    add_heading(doc, "Predicting Long NYC Yellow Taxi Trips Using Python-Based AI Analytics", level=1)

    add_heading(doc, "Introduction and Problem Framing", level=1)
    add_paragraph(
        doc,
        "New York City taxi service is a high-volume transportation system in which trip duration affects driver availability, passenger expectations, airport queues, congestion monitoring, and shift planning. This project framed a practical analytics question for taxi operators and city transportation analysts: Can a trip be classified as likely to last at least 30 minutes using structured trip information available in the NYC Taxi and Limousine Commission trip records? The project treats this as a business and societal problem because long trips influence how quickly vehicles return to dense pickup areas and how accurately riders and dispatch teams can plan service.",
    )
    add_paragraph(
        doc,
        "The selected artificial intelligence task was binary classification. The target variable, long_trip, equals 1 when a cleaned trip duration is at least 30 minutes and equals 0 otherwise. This target is easy for stakeholders to understand and supports operational decision-making. The model is not intended to deny rides, penalize drivers, or deprioritize neighborhoods. Instead, it is designed as a decision-support signal that can help teams anticipate vehicle availability and identify time periods or trip types that may create operational strain.",
    )

    add_heading(doc, "Dataset Exploration and Preprocessing", level=1)
    add_paragraph(
        doc,
        f"The analysis used the official NYC Taxi and Limousine Commission yellow taxi trip record dataset for January 2024 (New York City Taxi and Limousine Commission [NYC TLC], 2024a). The raw parquet file contained {metrics['dataset']['raw_rows']:,} records. The Python pipeline selected relevant fields, including pickup and dropoff timestamps, passenger count, trip distance, rate code, pickup and dropoff location IDs, fare fields, and total amount. The pipeline was built in Python with pandas, NumPy, Matplotlib, Seaborn, and scikit-learn so the work could run locally and in Google Colab (Hunter, 2007; McKinney, 2010; Pedregosa et al., 2011; Waskom, 2021).",
    )
    add_paragraph(
        doc,
        f"Preprocessing began by converting timestamp columns into datetime values and deriving trip_duration_min from the difference between dropoff and pickup times. The pipeline then applied a sequence of cleaning rules to remove records that likely reflected data-entry issues or trips outside the intended analysis scope. The script retained January 2024 pickups, durations between 1 and 180 minutes, distances between 0.1 and 60 miles, passenger counts from 1 to 6, positive fares and totals, standard TLC rate codes from 1 to 6, and valid pickup and dropoff zone IDs from 1 to 265. After cleaning, {metrics['dataset']['clean_rows_before_sampling']:,} realistic trips remained, representing 90.55% of the raw file.",
    )
    add_paragraph(
        doc,
        f"Because the full cleaned dataset was large, the model used a deterministic {metrics['dataset']['model_sample_rows']:,}-row sample with random_state = 42. This choice preserved reproducibility while keeping runtime reasonable in Colab. The sampled data had a median duration of {metrics['dataset']['median_duration_min']} minutes and a median distance of {metrics['dataset']['median_distance_miles']} miles. Long trips were a minority class: {metrics['target_distribution']['long_30_plus_min']:,} trips, or {metrics['target_distribution']['long_trip_rate'] * 100:.2f}%, lasted 30 minutes or longer.",
    )

    add_table_title(doc, 1, "Preprocessing Audit Summary")
    audit_rows = []
    selected_steps = {
        "03_filter_pickup_month",
        "04_filter_duration",
        "05_filter_distance",
        "06_filter_passengers",
        "07_filter_positive_fare",
        "09_filter_standard_ratecode",
        "12_feature_engineering",
        "13_model_sample",
    }
    for row in audit:
        if row["step"] in selected_steps:
            audit_rows.append(
                [
                    row["description"],
                    f"{int(row['rows_removed']):,}",
                    f"{int(row['rows_after']):,}",
                ]
            )
    add_apa_table(
        doc,
        ["Preprocessing step", "Rows removed", "Rows after"],
        audit_rows,
        [4.2, 1.1, 1.1],
    )

    add_paragraph(
        doc,
        "Feature engineering translated the raw trip records into inputs suitable for machine learning. The final numeric inputs were trip_distance, passenger_count, pickup_hour, pickup_dayofweek, is_weekend, and airport_trip. The categorical inputs were VendorID, RatecodeID, PULocationID, and DOLocationID. Location and rate-code identifiers were converted to strings so the model would treat them as categories rather than ordered quantities. The target was created only after deriving trip duration, and trip_duration_min was not used as a model input to avoid target leakage.",
    )

    add_table_title(doc, 2, "Feature Engineering Summary")
    feature_rows = [
        [row["column"], row["type"], row["purpose"]]
        for row in features
        if row["column"] in ("pickup_hour", "pickup_dayofweek", "is_weekend", "airport_trip", "long_trip")
    ]
    add_apa_table(
        doc,
        ["Column", "Type", "Purpose"],
        feature_rows,
        [1.5, 1.3, 3.6],
    )

    add_figure(
        doc,
        1,
        "Cleaned Trip Duration Distribution",
        REPORTS / "figures" / "duration_distribution.png",
        width=5.6,
    )

    add_heading(doc, "Initial Model Development", level=1)
    add_paragraph(
        doc,
        "The initial model was a class-balanced logistic regression classifier. Logistic regression was selected because it is fast, reproducible, and explainable enough for an early analytics deliverable. This was preferable for the first project stage because stakeholders can review the direction and relative strength of coefficients instead of receiving only black-box predictions. The pipeline used an 80/20 stratified train-test split, producing 96,000 training rows and 24,000 test rows while preserving the long-trip rate in both sets.",
    )
    add_paragraph(
        doc,
        "The scikit-learn pipeline used a ColumnTransformer to apply separate preprocessing to numeric and categorical columns. Numeric variables were median-imputed and standardized. Categorical variables were imputed with the most frequent value and one-hot encoded with handle_unknown = ignore, which allows the model to process future trips that contain categories not observed during training. A most-frequent dummy classifier served as the baseline because long trips are rare.",
    )

    add_table_title(doc, 3, "Held-Out Test-Set Model Performance")
    base = metrics["baseline_most_frequent"]
    model = metrics["logistic_regression"]
    add_apa_table(
        doc,
        ["Model", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC"],
        [
            [
                "Most frequent baseline",
                f"{base['accuracy']:.4f}",
                f"{base['precision']:.4f}",
                f"{base['recall']:.4f}",
                f"{base['f1']:.4f}",
                f"{base['roc_auc']:.4f}",
            ],
            [
                "Balanced logistic regression",
                f"{model['accuracy']:.4f}",
                f"{model['precision']:.4f}",
                f"{model['recall']:.4f}",
                f"{model['f1']:.4f}",
                f"{model['roc_auc']:.4f}",
            ],
        ],
        [2.1, 0.85, 0.85, 0.85, 0.75, 0.85],
    )
    add_paragraph(
        doc,
        f"The baseline reached {base['accuracy'] * 100:.2f}% accuracy because most trips were under 30 minutes, but its recall and F1 score for long trips were both 0. The logistic regression model improved the practical objective by reaching {model['recall'] * 100:.2f}% recall and a ROC-AUC of {model['roc_auc']:.4f}. This means the model identified most long trips in the test set, even though precision was {model['precision'] * 100:.2f}%. For an early warning tool, higher recall is valuable because missing long trips would weaken planning decisions. However, the precision result also shows the model should not be used for high-stakes automated decisions.",
    )
    add_paragraph(
        doc,
        "The coefficient output showed that trip distance had the strongest positive relationship with long-trip risk. Pickup and dropoff zones were also informative, which is useful operationally but ethically sensitive because location can act as a proxy for neighborhood, income, airport access, or commuting patterns. The model should therefore support aggregate planning and human review rather than individual ride denial, driver penalties, or neighborhood-level service restrictions.",
    )
    add_figure(
        doc,
        2,
        "Confusion Matrix for the Balanced Logistic Regression Model",
        REPORTS / "figures" / "confusion_matrix.png",
        width=4.8,
    )

    add_heading(doc, "Team Collaboration Process", level=1)
    add_paragraph(
        doc,
        "The project was organized as a small analytics team workflow. The major roles were stakeholder framing, data engineering and preprocessing, model development, and report quality assurance. GitHub was used for version control, with the completed work merged into the public repository on the main branch. Google Colab was configured as a collaboration environment so teammates could clone the repository, install requirements, run the pipeline, and inspect outputs directly in the notebook.",
    )
    add_table_title(doc, 4, "Team Member Contributions")
    add_apa_table(
        doc,
        ["Team member", "Primary contribution", "Specific project work"],
        [
            [
                "Soumith Kumar Ananthula",
                "Project coordination and stakeholder framing",
                "Defined the transportation problem, reviewed deliverables, coordinated GitHub/Colab workflow, and connected model outputs to the assignment requirements.",
            ],
            [
                "Varun Reddy Beesam",
                "Dataset exploration and preprocessing",
                "Reviewed TLC fields, documented cleaning rules, checked row-count changes, and helped verify preprocessing audit outputs.",
            ],
            [
                "Vivek Doppalapudi",
                "Model development and evaluation",
                "Supported model selection, train-test evaluation, metric interpretation, coefficient review, and ethical limitations for model use.",
            ],
        ],
        [1.7, 1.8, 2.9],
    )
    add_paragraph(
        doc,
        "The work followed an iterative analytics lifecycle. First, the team selected NYC taxi data because it matched the transportation dataset option and provided enough size and complexity for meaningful preprocessing. Second, the pipeline was built to download the official parquet file rather than relying on manually uploaded data. Third, model outputs and coefficients were reviewed, which led to a stronger preprocessing rule that retained only standard TLC rate codes. Fourth, the notebook was improved to show preprocessing steps, feature changes, target distribution, metrics, and charts inline so a reviewer can understand how the data changed before modeling.",
    )
    add_paragraph(
        doc,
        "The team also reflected on ethical use. The model does not use names or exact addresses, but zone-level transportation data can still reveal social and geographic patterns. For that reason, the recommended use is aggregate operational planning, not service exclusion or price manipulation. Future iterations should validate the model across multiple months, add weather and event context, compare additional algorithms, calibrate probabilities, and audit error rates by pickup and dropoff zone.",
    )

    doc.add_page_break()
    references = [
        "Hunter, J. D. (2007). Matplotlib: A 2D graphics environment. Computing in Science & Engineering, 9(3), 90-95. https://doi.org/10.1109/MCSE.2007.55",
        "McKinney, W. (2010). Data structures for statistical computing in Python. In S. van der Walt & J. Millman (Eds.), Proceedings of the 9th Python in Science Conference (pp. 56-61). https://doi.org/10.25080/Majora-92bf1922-00a",
        "New York City Taxi and Limousine Commission. (2024a). TLC trip record data. City of New York. https://www.nyc.gov/site/tlc/about/tlc-trip-record-data.page",
        "New York City Taxi and Limousine Commission. (2024b). Yellow taxi trip records data dictionary. City of New York. https://www.nyc.gov/assets/tlc/downloads/pdf/data_dictionary_trip_records_yellow.pdf",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, M., & Duchesnay, E. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825-2830. https://jmlr.org/papers/v12/pedregosa11a.html",
        "Waskom, M. L. (2021). Seaborn: Statistical data visualization. Journal of Open Source Software, 6(60), Article 3021. https://doi.org/10.21105/joss.03021",
    ]
    add_references(doc, references)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
